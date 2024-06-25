import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from encrypt_decrypt import*
from keys import*
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


color_map = {
        WD_COLOR_INDEX.YELLOW: "FFFF00",
        WD_COLOR_INDEX.BRIGHT_GREEN: "00FF00",
        WD_COLOR_INDEX.TURQUOISE: "00FFFF",
        WD_COLOR_INDEX.PINK: "FF00FF",
        WD_COLOR_INDEX.BLUE: "0000FF",
        WD_COLOR_INDEX.RED: "FF0000",
        WD_COLOR_INDEX.DARK_BLUE: "00008B",
        WD_COLOR_INDEX.TEAL: "008080",
        WD_COLOR_INDEX.GREEN: "008000",
        WD_COLOR_INDEX.VIOLET: "800080",
        WD_COLOR_INDEX.DARK_RED: "8B0000",
        WD_COLOR_INDEX.DARK_YELLOW: "FFD700",  # Dark yellow approximation
        WD_COLOR_INDEX.GRAY_25: "C0C0C0",
        WD_COLOR_INDEX.GRAY_50: "808080",
        WD_COLOR_INDEX.BLACK: "000000",
        WD_COLOR_INDEX.AUTO: "FFFFFF"
    }

def rgb_to_hex(rgb):
    return '#%02x%02x%02x' % (rgb[0], rgb[1], rgb[2])

def get_highlight_color(index):
    return color_map.get(index, "FFFFFF")

def get_color_index(bg_color):
    color_map_reverse = {v: k for k, v in color_map.items()}  # Tạo một từ điển ngược với mã màu là khóa và chỉ mục màu là giá trị

    # Lặp qua từ điển ngược để tìm mã màu tương ứng với bg_color
    for color_hex, color_index in color_map_reverse.items():
        if bg_color == color_hex:
            return color_index
    
    # Trả về màu tự động nếu không tìm thấy một màu nào khớp
    return WD_COLOR_INDEX.AUTO


def read_docx(file_path):
    doc = Document(file_path)
    character_info = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            color_hex = None
            bg_color_hex = None
            italic = 0
            size = None
            if run.font.color and run.font.color.rgb:
                color = run.font.color.rgb
                color_hex = '{:02x}{:02x}{:02x}'.format(color[0], color[1], color[2])

            if run.font.highlight_color:
                highlight_color = run.font.highlight_color
                if highlight_color in WD_COLOR_INDEX.__members__.values():
                    bg_color_hex = get_highlight_color(highlight_color)
            
            if run.font.italic:
                italic = 1

            size = None
            if run.font.size:
                size = run.font.size.pt

            for char in text:
                character_info.append({
                    'text': char,
                    'color': color_hex,
                    'bg_color': bg_color_hex,
                    'italic': italic,
                    'size': size
                })

    text = ''
    color = ''
    bg_color = ''
    italic = ''
    size = ''
    for info in character_info:
        text += info['text']
        color += info['color'] if info['color'] else '000000'
        bg_color += info['bg_color'] if info['bg_color'] else 'ffffff'
        italic += str(info['italic'])
        size += str(info['size']) + ',' if info['size'] else '10,'

    return text, color, bg_color, size[:-1], italic

def hex_to_rgb(hex_color):
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

def encrypt_message():
    k = key()
    
    text, color, bg_color, size , italic= '', '', '', '', ''
    
    current_index = '1.0'
    while True:
        next_index = entry_plain_text.index(f"{current_index} +1c")
        if next_index == current_index:
            break

        char = entry_plain_text.get(current_index)
        tag_names = entry_plain_text.tag_names(current_index)
        
        if tag_names:
            for tag in tag_names:
                if '_' in tag:
                    c, bg, s, it = tag.split('_')
                    text += char
                    color += c[1:]  # Remove the leading '#'
                    bg_color += bg[1:]  # Remove the leading '#'
                    size += s + ','
                    italic += it 
                    print(italic, '\n')
                else:
                    text += char
        elif color:
            text += char
            color += '000000'
            bg_color += 'ffffff'
            size += '10,'
            italic += '0'  # Default italic
        else:
            text += char
        current_index = next_index

    data = text
    if color:
        data += f'#{color}#{bg_color}#{size}#{italic}'
        print("this is input", data)
    encrypted_text = encrypt(data, k)

    text_encrypted.config(state=tk.NORMAL)
    text_encrypted.delete('1.0', tk.END)
    text_encrypted.insert(tk.END, encrypted_text)
    text_encrypted.config(state=tk.DISABLED)

def decrypt_message():
    encrypted_text = entry_encrypted_text_decrypt.get('1.0', tk.END).strip()
    data = encrypted_text
    if encrypted_text.endswith('.txt') or encrypted_text.endswith('.docx'):
        with open(encrypted_text, 'r', encoding='utf-8') as file:
            data = file.read()
   
    decrypted_data = cripher(data, key())
    print("day là ban ma: ", decrypted_data)
    parts = decrypted_data.split('#')
    text_data = parts[0]
    color_data = parts[1] if len(parts) > 1 else None
    bg_color_data = parts[2] if len(parts) > 2 else None
    size_data = parts[3][:-1].split(',') if len(parts) > 3 else None
    italic_data = parts[4] if len(parts) > 4 else None

    text_decrypted.config(state=tk.NORMAL)
    text_decrypted.delete('1.0', tk.END)
    if color_data:
        for i in range(len(text_data)):
            text = text_data[i]
            color = '#' + color_data[i*6:i*6+6]
            bg_color = '#' + bg_color_data[i*6:i*6+6]
            size = int(float(size_data[i]))  # Convert from float to int
            italic = int(italic_data[i])
            
            # Determine font style based on italic indicator
            font_style = "italic" if italic else "normal"
            
            tag = f"{color}_{bg_color}_{size}_{italic}"
            text_decrypted.insert(tk.END, text, tag)
            text_decrypted.tag_configure(tag, foreground=color, background=bg_color, font=("Helvetica", size, font_style))

    else:
        text_decrypted.insert(tk.END, decrypted_data)

    text_decrypted.config(state=tk.DISABLED)

def chuyen_text():
    encrypted_text = text_encrypted.get('1.0', tk.END).strip()
    entry_encrypted_text_decrypt.delete('1.0', tk.END)
    entry_encrypted_text_decrypt.insert('1.0', encrypted_text)

def browse_file_encryption():
    filename = filedialog.askopenfilename(
        initialdir=".",
        title="Select File",
        filetypes=(
            ("All files", "*.*"),
            ("Text files", "*.txt"), 
            ("Word files", "*.docx")
        )
    )

    data, color_data, bg_color_data, size_data, italic_data = '', '', '', '', ''
    if filename.endswith('.txt'):
        with open(filename, 'r', encoding='utf-8') as file:
            data = file.read()
    elif filename.endswith('.docx'):
        data, color_data, bg_color_data, size_data, italic_data  = read_docx(filename)
        size_data = size_data.split(',')

    entry_plain_text.delete('1.0', tk.END)
    if color_data:
        for i in range(len(data)):
            text = data[i]
            color = '#' + color_data[i*6:i*6+6]
            bg_color = '#' + bg_color_data[i*6:i*6+6]
            size = int(float(size_data[i]))
            italic = int(italic_data[i])
            
            # Determine font style based on italic indicator
            font_style = "italic" if italic else "normal"

            tag = f"{color}_{bg_color}_{size}_{italic}"
            entry_plain_text.insert(tk.END, text, tag)
            entry_plain_text.tag_configure(tag, foreground=color, background=bg_color, font=("Helvetica", size, font_style))
    else:
        entry_plain_text.insert(tk.END, data)

def browse_file_decryption():
    filename = filedialog.askopenfilename(initialdir="/", title="Select File", filetypes=(("Text files", "*.txt"), ("Word files", "*.docx"), ("All files", "*.*")))
    data= ''
    if filename:
        entry_encrypted_text_decrypt.delete('1.0', tk.END)
        with open(filename, 'r', encoding='utf-8') as file:
            data = file.read()
        entry_encrypted_text_decrypt.insert('1.0', data)

def save_encrypted_text():
    encrypted_text = text_encrypted.get('1.0', tk.END).strip()
    filename = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=(("Text files", "*.txt"), ("Word files", "*.docx"), ("All files", "*.*")))
    if filename:
        with open(filename, "w", encoding='utf-8') as file:
            file.write(encrypted_text)

def save_decrypted_text():
    decrypted_text = text_decrypted.get('1.0', tk.END).strip()
    

    filename = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=(("Text files", "*.txt"), ("Word files", "*.docx"), ("All files", "*.*")))
    if not filename:
        return  # Người dùng đã hủy hoặc không chọn file

    
    if filename.endswith(".txt"):
        with open(filename, "w", encoding='utf-8') as file:
            file.write(decrypted_text)
    elif filename.endswith(".docx"):
        doc = Document()
        paragraph = doc.add_paragraph()

        current_index = '1.0'
        while True:
            next_index = text_decrypted.index(f"{current_index} +1c")
            if next_index == current_index:
                break

            char = text_decrypted.get(current_index)
            tag_names = text_decrypted.tag_names(current_index)

            if tag_names:
                for tag in tag_names:
                    if '_' in tag:
                        color, bg_color, size, italic = tag.split('_')
                        run = paragraph.add_run(char)
                        run.font.color.rgb = hex_to_rgb(color[1:])  # Remove the leading '#'
                        
                        run.font.highlight_color = get_color_index(bg_color[1:])  # Áp dụng màu nền
                        run.font.size = Pt(int(size))
                        run.font.italic = True if italic == '1' else False
                    else:
                        run = paragraph.add_run(char)
            else:
                run = paragraph.add_run(char)

            current_index = next_index

        doc.save(filename)
    

# Create the main window
root = tk.Tk()
root.title("Hệ mã hóa AES")

# Get the screen width and height
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Set the window size to 80% of the screen width and height
window_width = int(screen_width * 0.8)
window_height = int(screen_height * 0.8)

# Center the window on the screen
window_x = (screen_width - window_width) // 2
window_y = (screen_height - window_height) // 2

# Set the geometry of the window
root.geometry(f"{window_width}x{window_height}+{window_x}+{window_y}")

# Configure the grid to expand with window resizing
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=0)  # Separator column
root.grid_columnconfigure(2, weight=1)
root.grid_rowconfigure(0, weight=0)  # Top row for label
root.grid_rowconfigure(1, weight=1)  # Rest of the window

# Create frames
frame_encrypt = ttk.Frame(root)
frame_encrypt.grid(row=1, column=0, sticky="NSEW")

frame_decrypt = ttk.Frame(root)
frame_decrypt.grid(row=1, column=2, sticky="NSEW")

# Add a separator
separator = ttk.Separator(root, orient='vertical')
separator.grid(row=1, column=1, sticky='NS')

# Configure frames to expand with window resizing
frame_encrypt.grid_columnconfigure(0, weight=1)
frame_encrypt.grid_columnconfigure(1, weight=1)
frame_encrypt.grid_columnconfigure(2, weight=1)
frame_encrypt.grid_columnconfigure(3, weight=1)
frame_encrypt.grid_columnconfigure(4, weight=1)
frame_encrypt.grid_rowconfigure(0, weight=1)
frame_encrypt.grid_rowconfigure(1, weight=1)
frame_encrypt.grid_rowconfigure(2, weight=1)
frame_encrypt.grid_rowconfigure(3, weight=1)
frame_encrypt.grid_rowconfigure(4, weight=1)


frame_decrypt.grid_columnconfigure(0, weight=1)
frame_decrypt.grid_columnconfigure(1, weight=1)
frame_decrypt.grid_columnconfigure(2, weight=1)
frame_decrypt.grid_columnconfigure(3, weight=1)
frame_decrypt.grid_columnconfigure(4, weight=1)
frame_decrypt.grid_rowconfigure(0, weight=1)
frame_decrypt.grid_rowconfigure(1, weight=1)
frame_decrypt.grid_rowconfigure(2, weight=1)
frame_decrypt.grid_rowconfigure(3, weight=1)
frame_decrypt.grid_rowconfigure(4, weight=1)

# Encryption section
ttk.Label(root, text="Hệ mã hóa AES", font=("Arial", 20)).grid(row=0, column=0, columnspan=3, pady=10)

ttk.Label(frame_encrypt, text="MÃ HÓA", font=("Arial", 16)).grid(row=0, column=0, columnspan=4, pady=10)

ttk.Label(frame_encrypt, text="Bản rõ:", font=("Arial", 12)).grid(row=1, column=0, columnspan=1, sticky="E")
entry_plain_text = tk.Text(frame_encrypt, width=45, height=5)
entry_plain_text.grid(row=1, column=1, padx=5, pady=5, sticky="W")
tk.Button(frame_encrypt, text="Chọn file", command=browse_file_encryption, bg='lightblue', width=10, height=2).grid(row=1, column=2, columnspan=2, padx=5, pady=5)

tk.Button(frame_encrypt, text="Mã hóa", font=("Arial", 12), command=encrypt_message, bg='lightblue', width=10, height=2).grid(row=2, column=0, columnspan=4, pady=5)

ttk.Label(frame_encrypt, text="Bản mã:", font=("Arial", 12)).grid(row=3, column=0, columnspan=1, sticky="E")
text_encrypted = tk.Text(frame_encrypt, width=45, height=5, state=tk.DISABLED)
text_encrypted.grid(row=3, column=1, sticky="W")

tk.Button(frame_encrypt, text="Chuyển", command=chuyen_text, bg='lightblue', width=10, height=2).grid(row=3, column=2, padx=5, pady=2)  # Adjust pady for smaller spacing
tk.Button(frame_encrypt, text="Lưu", command=save_encrypted_text, bg='lightblue', width=10, height=2).grid(row=4, column=2, padx=5, pady=2)  # Move this button to the next row with smaller spacing
# Decryption section
ttk.Label(frame_decrypt, text="GIẢI MÃ", font=("Arial", 16)).grid(row=0, column=0, columnspan=4, pady=10)

ttk.Label(frame_decrypt, text="Bản mã:", font=("Arial", 12)).grid(row=1, column=0, sticky="E")
entry_encrypted_text_decrypt = tk.Text(frame_decrypt, width=45, height=5)
entry_encrypted_text_decrypt.grid(row=1, column=1, padx=5, pady=5, sticky="W")
tk.Button(frame_decrypt, text="Chọn file", command=browse_file_decryption, bg='lightblue', width=10, height=2).grid(row=1, column=3,columnspan=2, padx=5, pady=5)

tk.Button(frame_decrypt, text="Giải mã", font=("Arial", 12), command=decrypt_message, bg='lightblue', width=10, height=2).grid(row=2, column=1, columnspan=2, pady=5)

ttk.Label(frame_decrypt, text="Bản rõ:", font=("Arial", 12)).grid(row=3, column=0, sticky="E")
text_decrypted = tk.Text(frame_decrypt, width=45, height=5, state=tk.DISABLED)
text_decrypted.grid(row=3, column=1, padx=5, pady=5, sticky="W")

tk.Button(frame_decrypt, text="Lưu", command=save_decrypted_text, bg='lightblue', width=10, height=2).grid(row=3, column=2, columnspan=2, pady=5)

# Add some space around the elements
for child in frame_encrypt.winfo_children(): 
    child.grid_configure(padx=5, pady=5)

for child in frame_decrypt.winfo_children(): 
    child.grid_configure(padx=5, pady=5)

# Run the application
root.mainloop()

